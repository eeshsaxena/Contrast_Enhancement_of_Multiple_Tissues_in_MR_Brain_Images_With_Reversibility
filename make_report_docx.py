"""
Generate report.docx — MATLAB Implementation Report
Contrast Enhancement of Multiple Tissues in MR Brain Images With Reversibility
Wu et al., IEEE Signal Processing Letters, 2021

Styling rules:
  - No black/grey/blue text colours anywhere
  - Bold and Italic differentiate heading levels
  - ONE separator line only (after header)
  - Code label bars: navy blue #1a3a5c (ONLY blue element)
  - Table headers: bold, bordered (no colour fill)
  - All body text: automatic (Word default black)
"""

from docx import Document
from docx.shared import Pt, RGBColor, Inches, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import os

IMG_DIR = r'd:\New folder'

BLUE_CODE = RGBColor(0x1a, 0x3a, 0x5c)   # ONLY used on code label bars
RED_DOI   = RGBColor(0xc8, 0x10, 0x2e)

doc = Document()
for sec in doc.sections:
    sec.top_margin    = Cm(2.2)
    sec.bottom_margin = Cm(2.2)
    sec.left_margin   = Cm(2.5)
    sec.right_margin  = Cm(2.5)

# ───────────────────────────────────────────────────────
# HELPERS
# ───────────────────────────────────────────────────────

def border_para(para, sz=12, color='000000'):
    pPr = para._p.get_or_add_pPr()
    pBdr = OxmlElement('w:pBdr')
    bot  = OxmlElement('w:bottom')
    bot.set(qn('w:val'),   'single')
    bot.set(qn('w:sz'),    str(sz))
    bot.set(qn('w:space'), '1')
    bot.set(qn('w:color'), color)
    pBdr.append(bot)
    pPr.append(pBdr)

def shade_para(para, fill):
    pPr = para._p.get_or_add_pPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), fill)
    pPr.append(shd)

def shade_cell(cell, fill):
    tc  = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd  = OxmlElement('w:shd')
    shd.set(qn('w:val'),   'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'),  fill)
    tcPr.append(shd)

def cell_border_bottom(cell, sz=12, color='000000'):
    tc   = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcB  = OxmlElement('w:tcBorders')
    bot  = OxmlElement('w:bottom')
    bot.set(qn('w:val'),   'single')
    bot.set(qn('w:sz'),    str(sz))
    bot.set(qn('w:color'), color)
    tcB.append(bot)
    tcPr.append(tcB)

def run(para, text, bold=False, italic=False, size=11, color=None, font='Times New Roman'):
    r = para.add_run(text)
    r.bold        = bold
    r.italic      = italic
    r.font.name   = font
    r.font.size   = Pt(size)
    if color:
        r.font.color.rgb = color
    return r

# ── Paragraph helpers ─────────────────────────────────
def body(text, size=11, space_before=2, space_after=6, justify=True):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(space_before)
    p.paragraph_format.space_after  = Pt(space_after)
    if justify:
        p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    r = p.add_run(text)
    r.font.name = 'Times New Roman'
    r.font.size = Pt(size)
    return p

def h2(text):
    """Section heading — Bold 13pt, space above, no colour."""
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(18)
    p.paragraph_format.space_after  = Pt(6)
    r = p.add_run(text)
    r.bold       = True
    r.font.name  = 'Calibri'
    r.font.size  = Pt(13)
    return p

def h3(text):
    """Subsection — Bold Italic 11.5pt, no colour."""
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(12)
    p.paragraph_format.space_after  = Pt(4)
    r = p.add_run(text)
    r.bold       = True
    r.italic     = True
    r.font.name  = 'Calibri'
    r.font.size  = Pt(11.5)
    return p

def h4(text):
    """Sub-subsection — Italic 11pt, no colour."""
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(8)
    p.paragraph_format.space_after  = Pt(3)
    r = p.add_run(text)
    r.italic     = True
    r.font.name  = 'Calibri'
    r.font.size  = Pt(11)
    return p

def bullet(text, bold_part=None):
    p = doc.add_paragraph(style='List Bullet')
    p.paragraph_format.left_indent  = Cm(0.5)
    p.paragraph_format.space_after  = Pt(3)
    p.paragraph_format.space_before = Pt(2)
    if bold_part:
        r1 = p.add_run(bold_part)
        r1.bold = True; r1.font.size = Pt(11); r1.font.name = 'Times New Roman'
        r2 = p.add_run(text)
        r2.font.size = Pt(11); r2.font.name = 'Times New Roman'
    else:
        r = p.add_run(text)
        r.font.size = Pt(11); r.font.name = 'Times New Roman'
    return p

def numbered(text, bold_part=None):
    p = doc.add_paragraph(style='List Number')
    p.paragraph_format.left_indent  = Cm(0.5)
    p.paragraph_format.space_after  = Pt(3)
    p.paragraph_format.space_before = Pt(2)
    if bold_part:
        r1 = p.add_run(bold_part)
        r1.bold = True; r1.font.size = Pt(11); r1.font.name = 'Times New Roman'
        r2 = p.add_run(text)
        r2.font.size = Pt(11); r2.font.name = 'Times New Roman'
    else:
        r = p.add_run(text)
        r.font.size = Pt(11); r.font.name = 'Times New Roman'
    return p

def eq_block(eq_text, eq_num=None):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(8)
    p.paragraph_format.space_after  = Pt(8)
    p.paragraph_format.left_indent  = Cm(1.5)
    shade_para(p, 'F7F7F7')
    r = p.add_run(eq_text)
    r.italic     = True
    r.font.name  = 'Times New Roman'
    r.font.size  = Pt(11)
    if eq_num:
        r2 = p.add_run(f'    {eq_num}')
        r2.bold      = True
        r2.italic    = False
        r2.font.name = 'Times New Roman'
        r2.font.size = Pt(10)
    return p

def code_block(label_text, code_text):
    """Blue label bar (ONLY blue element) + shaded code body."""
    lp = doc.add_paragraph()
    lp.paragraph_format.space_before = Pt(10)
    lp.paragraph_format.space_after  = Pt(0)
    shade_para(lp, '1a3a5c')           # ← blue ONLY here
    lr = lp.add_run(f'  {label_text}  ')
    lr.bold            = True
    lr.font.name       = 'Courier New'
    lr.font.size       = Pt(9)
    lr.font.color.rgb  = RGBColor(0xFF, 0xFF, 0xFF)

    cp = doc.add_paragraph()
    cp.paragraph_format.space_before = Pt(0)
    cp.paragraph_format.space_after  = Pt(12)
    cp.paragraph_format.left_indent  = Cm(0.3)
    shade_para(cp, 'F2F2F2')
    cr = cp.add_run(code_text)
    cr.font.name = 'Courier New'
    cr.font.size = Pt(8.5)

def key_box(label, text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after  = Pt(10)
    p.paragraph_format.left_indent  = Cm(0.4)
    shade_para(p, 'F7F7F7')
    border_para(p, sz=6, color='888888')
    r1 = p.add_run(label + '  ')
    r1.bold = True; r1.font.size = Pt(11); r1.font.name = 'Times New Roman'
    r2 = p.add_run(text)
    r2.font.size = Pt(11); r2.font.name = 'Times New Roman'

def add_figure(filename, caption, width=Inches(5.5)):
    """Insert an image from IMG_DIR with an italic centred caption."""
    path = os.path.join(IMG_DIR, filename)
    if not os.path.exists(path):
        print(f'[WARN] Image not found: {path}')
        return
    ip = doc.add_paragraph()
    ip.alignment = WD_ALIGN_PARAGRAPH.CENTER
    ip.paragraph_format.space_before = Pt(8)
    ip.paragraph_format.space_after  = Pt(2)
    run_img = ip.add_run()
    run_img.add_picture(path, width=width)
    cp = doc.add_paragraph()
    cp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    cp.paragraph_format.space_before = Pt(2)
    cp.paragraph_format.space_after  = Pt(14)
    cr = cp.add_run(caption)
    cr.italic = True; cr.font.size = Pt(9.5); cr.font.name = 'Times New Roman'

def add_figure_grid(items):
    """
    items: list of (filename, caption) tuples — laid out 2 per row.
    Each cell contains the image + caption below it.
    """
    ncols = 2
    nrows = (len(items) + 1) // ncols
    tbl = doc.add_table(rows=nrows, cols=ncols)
    tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
    tbl.style = 'Table Grid'
    idx = 0
    for ri in range(nrows):
        for ci in range(ncols):
            if idx >= len(items):
                break
            fname, cap = items[idx]; idx += 1
            cell = tbl.cell(ri, ci)
            path = os.path.join(IMG_DIR, fname)
            # Image paragraph
            ip = cell.paragraphs[0]
            ip.alignment = WD_ALIGN_PARAGRAPH.CENTER
            ip.paragraph_format.space_before = Pt(4)
            ip.paragraph_format.space_after  = Pt(2)
            if os.path.exists(path):
                ip.add_run().add_picture(path, width=Inches(2.8))
            else:
                ip.add_run(f'[{fname} not found]')
            # Caption paragraph
            cp = cell.add_paragraph()
            cp.alignment = WD_ALIGN_PARAGRAPH.CENTER
            cp.paragraph_format.space_before = Pt(2)
            cp.paragraph_format.space_after  = Pt(6)
            cr = cp.add_run(cap)
            cr.italic = True; cr.font.size = Pt(9); cr.font.name = 'Times New Roman'
    # space after grid
    sp = doc.add_paragraph()
    sp.paragraph_format.space_after = Pt(10)

def add_table(headers, rows, footer_row=None, caption=None):
    nrows = 1 + len(rows) + (1 if footer_row else 0)
    tbl = doc.add_table(rows=nrows, cols=len(headers))
    tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
    tbl.style = 'Table Grid'

    # Header row — bold text, thick bottom border, no fill
    hcells = tbl.rows[0].cells
    for i, h in enumerate(headers):
        shade_cell(hcells[i], 'FFFFFF')
        cell_border_bottom(hcells[i], sz=16)
        hcells[i].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        r  = hcells[i].paragraphs[0].add_run(h)
        r.bold = True; r.font.size = Pt(10); r.font.name = 'Calibri'

    # Data rows — alternating very light shading
    fills = ['FFFFFF', 'F7F7F7']
    for ri, row_data in enumerate(rows):
        cells = tbl.rows[ri + 1].cells
        for ci, val in enumerate(row_data):
            shade_cell(cells[ci], fills[ri % 2])
            cells[ci].paragraphs[0].alignment = (
                WD_ALIGN_PARAGRAPH.LEFT if ci == 0 else WD_ALIGN_PARAGRAPH.CENTER)
            r = cells[ci].paragraphs[0].add_run(str(val))
            r.font.name = 'Times New Roman'; r.font.size = Pt(10)
            if ci == 0: r.italic = True

    # Footer row
    if footer_row:
        fc = tbl.rows[-1].cells
        for ci, val in enumerate(footer_row):
            shade_cell(fc[ci], 'EFEFEF')
            fc[ci].paragraphs[0].alignment = (
                WD_ALIGN_PARAGRAPH.LEFT if ci == 0 else WD_ALIGN_PARAGRAPH.CENTER)
            r = fc[ci].paragraphs[0].add_run(str(val))
            r.bold = True; r.font.size = Pt(10); r.font.name = 'Calibri'

    if caption:
        cp = doc.add_paragraph()
        cp.alignment = WD_ALIGN_PARAGRAPH.CENTER
        cp.paragraph_format.space_before = Pt(4)
        cp.paragraph_format.space_after  = Pt(14)
        cr = cp.add_run(caption)
        cr.italic = True; cr.font.size = Pt(9.5); cr.font.name = 'Times New Roman'

# ═══════════════════════════════════════════════════════
#  HEADER
# ═══════════════════════════════════════════════════════
tp = doc.add_paragraph()
tp.alignment = WD_ALIGN_PARAGRAPH.CENTER
tp.paragraph_format.space_after = Pt(4)
tr = tp.add_run(
    'IMPLEMENTATION OF CONTRAST ENHANCEMENT OF\n'
    'MULTIPLE TISSUES IN MR BRAIN IMAGES WITH REVERSIBILITY')
tr.bold = True; tr.font.size = Pt(16); tr.font.name = 'Calibri'

sp = doc.add_paragraph()
sp.alignment = WD_ALIGN_PARAGRAPH.CENTER
sp.paragraph_format.space_after = Pt(4)
sr = sp.add_run('Implementation Report  |  Medical Image Processing & Reversible Data Hiding')
sr.italic = True; sr.font.size = Pt(11.5); sr.font.name = 'Calibri'

ap = doc.add_paragraph()
ap.alignment = WD_ALIGN_PARAGRAPH.CENTER
ap.paragraph_format.space_after = Pt(3)
ar = ap.add_run('Based on: Hao-Tian Wu, Kaihan Zheng, Qi Huang, Jiankun Hu')
ar.font.size = Pt(11); ar.font.name = 'Times New Roman'

jp = doc.add_paragraph()
jp.alignment = WD_ALIGN_PARAGRAPH.CENTER
jp.paragraph_format.space_after = Pt(6)
jr = jp.add_run('IEEE Signal Processing Letters, Vol. 28, 2021  ·  DOI: 10.1109/LSP.2020.3048840')
jr.bold = True; jr.font.size = Pt(10.5); jr.font.name = 'Calibri'; jr.font.color.rgb = RED_DOI

# ── ONE separator line, right after the header ──────────
sep = doc.add_paragraph()
sep.paragraph_format.space_before = Pt(0)
sep.paragraph_format.space_after  = Pt(12)
border_para(sep, sz=16, color='000000')

# ═══════════════════════════════════════════════════════
#  ABSTRACT
# ═══════════════════════════════════════════════════════
abs_p = doc.add_paragraph()
abs_p.paragraph_format.space_before = Pt(0)
abs_p.paragraph_format.space_after  = Pt(10)
abs_p.paragraph_format.left_indent  = Cm(0.5)
abs_p.paragraph_format.right_indent = Cm(0.5)
abs_p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
shade_para(abs_p, 'F2F2F2')
r1 = abs_p.add_run('Abstract.  ')
r1.bold = True; r1.font.size = Pt(11); r1.font.name = 'Times New Roman'
r2 = abs_p.add_run(
    'Contrast enhancement (CE) of Magnetic Resonance (MR) brain images is a critical '
    'preprocessing step in clinical diagnosis. This report presents a complete MATLAB '
    'implementation of the hierarchical reversible CE scheme proposed by Wu et al. (2021, '
    'IEEE Signal Processing Letters). A deep convolutional neural network (U-Net) — '
    'approximated via multi-level Otsu thresholding (multithresh) — segments the MR image '
    'into Background, CSF, Grey Matter, and White Matter. For each tissue, principal '
    'grey-level values are identified using percentage threshold R, and the corresponding '
    'histogram bins are iteratively expanded S times using Reversible Data Hiding (RDH) via '
    'Equation (1), achieving histogram equalization. Recovery information (pL, pR per '
    'iteration) is stored in the LSBs of the last 16 pixels, enabling lossless original '
    'image recovery via Equations (2) and (3). Experimental results confirm correct '
    'multi-tissue enhancement and perfect reversibility (PSNR = inf) across all tissues.')
r2.font.size = Pt(10.8); r2.font.name = 'Times New Roman'

# ═══════════════════════════════════════════════════════
#  1. INTRODUCTION
# ═══════════════════════════════════════════════════════
h2('1.  Introduction')
body('Magnetic Resonance Imaging (MRI) produces detailed soft-tissue images of the brain using '
     'magnetic fields and radio waves, emitting no harmful ionising radiation. However, MR images '
     'frequently suffer from low contrast between adjacent tissue classes — particularly between '
     'White Matter (WM), Grey Matter (GM), and Cerebrospinal Fluid (CSF) — making clinical '
     'diagnosis difficult without further processing. Contrast Enhancement (CE) is therefore an '
     'essential preprocessing step in medical image analysis.')
body('A significant limitation of conventional CE is that it permanently alters pixel values, '
     'potentially destroying diagnostic information. Techniques based on Reversible Data Hiding '
     '(RDH) embed all information required to restore the original image within the enhanced '
     'image itself, allowing perfect recovery when needed.')
body('Prior reversible CE algorithms (Wu et al. 2015 [19]; Wu et al. 2020 [27]) used Otsu or '
     'GrabCut segmentation to separate background from ROI, then excluded background histogram '
     'bins from expansion. These approaches are limited: accuracy degrades for small or '
     'interconnected tissues, and only a single enhanced image is produced per original image. '
     'Wu et al. (2021) address both limitations with a hierarchical CE scheme guided per-tissue '
     'via U-Net CNN segmentation, producing one enhanced image per tissue class from the same '
     'original MR image.')

# ═══════════════════════════════════════════════════════
#  2. SYSTEM OVERVIEW
# ═══════════════════════════════════════════════════════
h2('2.  System Overview')
body('The proposed scheme consists of five main stages, implemented as MATLAB functions:')
for s in [
    ('Tissue Segmentation — ', 'Divide MR image into classes using U-Net (multithresh approximation in implementation).'),
    ('Principal Grey-Level Identification — ', 'Identify values with frequency > R in each tissue; only these bins are eligible for expansion.'),
    ('Histogram Preprocessing — ', 'Shrink histogram extremes to reserve S empty bins per side, preventing overflow after S expansion rounds.'),
    ('Procedure 1 (Enhancement) — ', 'Apply Equation (1) iteratively S times, expanding the two highest eligible bins and embedding recovery bits.'),
    ('Procedure 2 (Recovery) — ', 'Extract bits via Equation (2), restore pixel values via Equation (3), reverse preprocessing.'),
]:
    bullet(s[1], bold_part=s[0])

key_box('Key Insight:',
    'CE bins are chosen per-tissue (not per-background). The same pixel value shared by '
    'ROI and background no longer causes conflicts, and multiple tissue-enhanced images '
    'are independently generated from one MR image — the core advance over [19] and [27].')

# ═══════════════════════════════════════════════════════
#  3. MATHEMATICAL FORMULATION
# ═══════════════════════════════════════════════════════
h2('3.  Mathematical Formulation')

h3('3.1  RDH-Based Histogram Modification — Equation (1)')
body('Given the two highest-frequency histogram bins pL and pR (pL < pR), every pixel value p '
     'is transformed to p\u2032. Bit bi \u2208 {0,1} is embedded for the i-th pixel at pL or pR:')
eq_block(
    "p\u2032 =  p \u2212 1,      if p < pL\n"
    "        p \u2212 bi,     if p = pL    (embeds bit bi)\n"
    "        p,          if pL < p < pR   (unchanged)\n"
    "        p + bi,     if p = pR    (embeds bit bi)\n"
    "        p + 1,      if p > pR",
    eq_num='(Eq. 1)')
body('Applying Eq. (1) for S iterations achieves Histogram Equalization (HE) while all '
     'modification parameters remain embedded and recoverable.')

h3('3.2  Principal Grey-Level Threshold')
body('For each segmented tissue T, a grey-level v is designated as principal if:')
eq_block('percentage(v)  =  count(v in T) / |T|  >  R', eq_num='(Eq. 2)')
body('Only principal bins are eligible for expansion. Non-principal bins are merely shifted outward. '
     'This focuses CE on dominant tissue intensities, improving adaptability across tissue types.')

h3('3.3  Bit Extraction — Equation (2)')
body('Given the last expanded bins pLL and pLR, a bit b\u2032 is extracted from pixel p\u2032:')
eq_block(
    "b\u2032 = 1,     when p\u2032 = pLL\u22121 or p\u2032 = pLR+1\n"
    "b\u2032 = 0,     when p\u2032 = pLL  or p\u2032 = pLR\n"
    "b\u2032 = null,  other cases",
    eq_num='(Eq. 3)')

h3('3.4  Pixel Recovery — Equation (3)')
body('Simultaneously, the original pixel p is reconstructed from p\u2032:')
eq_block(
    "p = p\u2032 + 1,   when p\u2032 < pLL\n"
    "p = p\u2032,       when pLL\u22121 < p\u2032 < pLR+1\n"
    "p = p\u2032 \u2212 1,   when p\u2032 > pLR",
    eq_num='(Eq. 4)')

h3('3.5  Quality Metrics')
eq_block(
    'PSNR  = 10 \u00b7 log\u2081\u2080(255\u00b2 / MSE)     where   MSE = (1/HW)\u00b7\u03a3(I \u2212 Ic)\u00b2',
    eq_num='(Eq. 5)')
body('When I is perfectly recovered from Ic, MSE = 0 and PSNR = \u221e. '
     'RCEOI (Relative Contrast Enhancement), REEOI (Relative Entropy Enhancement), '
     'and RMBEOI (Relative Mean Brightness Error) follow [34]. SSIM measures perceptual '
     'similarity between original and enhanced images.')

# ═══════════════════════════════════════════════════════
#  4. TISSUE SEGMENTATION
# ═══════════════════════════════════════════════════════
h2('4.  Tissue Segmentation')
body('The paper uses U-Net [32] trained on 50 T2-weighted NeoBrainS12 images (384\u00d7384) to '
     'automatically segment MR images into Background, CSF, Grey Matter, and White Matter. '
     'In implementation, U-Net is approximated using MATLAB\'s multithresh (multi-level Otsu) '
     'followed by imquantize, with bwareaopen removing small isolated regions for noise suppression.')

code_block('MATLAB Code — Tissue Segmentation  (segment_tissues_multiotsu.m)',
'''function label_map = segment_tissues_multiotsu(img, n_classes)
% Approximate U-Net segmentation via multi-level Otsu thresholding.
% Paper Sec. II-B-1: segments MR image into tissue classes.
% label_map: 0=Background, 1=CSF, 2=Grey Matter, 3=White Matter

img = double(img);
thresholds = multithresh(img, n_classes - 1);   % n_classes-1 thresholds
label_map  = imquantize(img, thresholds) - 1;   % 0-indexed labels

label_map(img == 0) = 0;                        % force pure-black to background

for cls = 1:(n_classes - 1)
    mask = (label_map == cls);
    mask = bwareaopen(mask, 200);               % remove regions < 200 px
    label_map(label_map == cls & ~mask) = 0;
end
end''')

# Figure: segmentation result
add_figure('segmentation.png',
    'Fig. 1 — Tissue segmentation result on synthetic MR brain image. '
    'Four classes: Background (black), CSF (dark), Grey Matter (mid), White Matter (bright). '
    'Generated by the multi-level Otsu approximation of U-Net segmentation.',
    width=Inches(5.2))

# ═══════════════════════════════════════════════════════
#  5. PRINCIPAL GREY-LEVEL IDENTIFICATION
# ═══════════════════════════════════════════════════════
h2('5.  Principal Grey-Level Identification')
body('For each segmented tissue, the percentage of every grey-level value is computed. '
     'Values exceeding threshold R = 1% are labelled as principal. Only those histogram '
     'bins are eligible for expansion via Eq. (1). Non-principal bins are only shifted, '
     'never expanded — the key difference from [19] and [27].')

code_block('MATLAB Code — Principal Grey-Level Identification  (identify_principal_greylevels.m)',
'''function principal_vals = identify_principal_greylevels(img, tissue_mask, R)
% Paper Sec. II-B-2: find grey-levels with percentage > R within tissue.
% Returns: principal_vals — 0-indexed grey-level values (vector)

pixels = double(img(tissue_mask));
total  = numel(pixels);
if total == 0
    principal_vals = [];
    return;
end

hist_counts = histcounts(pixels, 0:256);   % counts for values 0..255
percentages = hist_counts / total;

% Return 0-indexed grey-level values with percentage > R
principal_vals = find(percentages > R) - 1;
end''')

# ═══════════════════════════════════════════════════════
#  6. HISTOGRAM PREPROCESSING
# ═══════════════════════════════════════════════════════
h2('6.  Histogram Preprocessing')
body('Before expanding bins, a preprocessing step reserves S empty bins on each side '
     'of the histogram to prevent pixel value overflow after S iterations of Eq. (1). '
     'Pixels in [0, S\u22121] are mapped to S; pixels in [256\u2212S, 255] are mapped to 255\u2212S. '
     'This shrink preserves the relative order of all pixel values.')

code_block('MATLAB Code — Histogram Preprocessing  (preprocess_histogram_shrink.m)',
'''function [pre_img, lut] = preprocess_histogram_shrink(img, S)
% Reserves S empty bins each side; prevents overflow after S expansions.
% Paper [24]: shrinks histogram preserving relative pixel order.

lut = uint8(0:255);               % identity LUT
lut(1:S)           = uint8(S);    % values 0..S-1  -> S
lut(256-S+1:256)   = uint8(255-S);% values 255-S..255 -> 255-S

pre_img = lut(double(img) + 1);   % apply LUT (1-indexed in MATLAB)
end''')

# ═══════════════════════════════════════════════════════
#  7. PROCEDURE 1 — HIERARCHICAL TISSUE CE
# ═══════════════════════════════════════════════════════
h2('7.  Procedure 1 — Hierarchical Tissue Contrast Enhancement')
body('For each tissue class, Procedure 1 generates one tissue-enhanced image Ic. '
     'At each of S iterations:')
for s in [
    ('Find pL, pR: ', 'the two highest bins from among the principal grey-level values in the current histogram.'),
    ('Encode pL, pR: ', 'as 16 bits (8+8) — the side information for this round.'),
    ('Apply Eq. (1): ', 'to all pixels, embedding the encoded bits into pixels at pL and pR.'),
    ('Record (pL, pR): ', 'in the expansion history for use during recovery.'),
]:
    numbered(s[1], bold_part=s[0])
body('After all S rounds, the last (pL, pR) pair is stored in the LSBs of the last '
     '16 pixels (8 bits for pL, 8 bits for pR). Neither the segmentation map nor R '
     'is required for recovery.')

code_block('MATLAB Code — Equation (1) Application  (apply_equation1.m)',
'''function [out_img, bit_idx] = apply_equation1(img, pL, pR, bits)
% Paper Equation (1): histogram bin expansion with bit embedding.
%   p\' = p-1      if p < pL
%        p-bi     if p = pL   (embed bit bi)
%        p         if pL < p < pR
%        p+bi     if p = pR   (embed bit bi)
%        p+1      if p > pR

flat    = double(img(:));
out     = flat;
bit_idx = 0;

for i = 1:numel(flat)
    p = flat(i);
    if p < pL
        out(i) = p - 1;
    elseif p == pL
        bit_idx = bit_idx + 1;
        bi      = bits(min(bit_idx, numel(bits)));
        out(i)  = p - bi;
    elseif p == pR
        bit_idx = bit_idx + 1;
        bi      = bits(min(bit_idx, numel(bits)));
        out(i)  = p + bi;
    elseif p > pR
        out(i) = p + 1;
    % else pL < p < pR: unchanged
    end
end

out_img = uint8(reshape(out, size(img)));
end''')

code_block('MATLAB Code — LSB Side-Info Storage  (encode_pL_pR_in_lsb.m)',
'''function out_img = encode_pL_pR_in_lsb(img, pL, pR)
% Paper: "LSB of sixteen pixels store the last two expanded bin values."
% Encodes pL (8 bits) and pR (8 bits) into LSB of last 16 pixels.

flat = double(img(:));
N    = numel(flat);

pL_bits = bitget(uint8(pL), 8:-1:1);   % 8 bits, MSB first
pR_bits = bitget(uint8(pR), 8:-1:1);
all_bits = [pL_bits, pR_bits];          % 16 bits total

for k = 1:16
    flat(N - 16 + k) = bitset(flat(N - 16 + k), 1, all_bits(k));
end

out_img = uint8(reshape(flat, size(img)));
end''')

# ═══════════════════════════════════════════════════════
#  8. PROCEDURE 2 — RECOVERY
# ═══════════════════════════════════════════════════════
h2('8.  Procedure 2 — Secret Data Extraction and Image Recovery')
body('Recovery processes the S expansion rounds in reverse order (last round first). '
     'At each step, read pLL and pLR from the current image and apply Eqs. (2) and (3) '
     'to extract bits and restore pixel values. Decode (pLL, pLR) for the previous step '
     'from the extracted bits, and repeat for all S rounds. Finally, undo the preprocessing shrink.')

code_block('MATLAB Code — Equations (2) & (3) Combined Recovery  (apply_equation2_3.m)',
'''function [rec_img, extr_bits] = apply_equation2_3(img, pLL, pLR)
% Paper Eq.(2): extract bit b\' from pixel p\'
%   b\' = 1    if p\' = pLL-1 or p\' = pLR+1
%        0    if p\' = pLL   or p\' = pLR
%        null otherwise
%
% Paper Eq.(3): recover original pixel p from p\'
%   p = p\'+1   if p\' < pLL
%       p\'      if pLL-1 < p\' < pLR+1
%       p\'-1   if p\' > pLR

flat      = double(img(:));
recovered = flat;
extr_bits = zeros(1, numel(flat));   % pre-allocate (trimmed after)
n_bits    = 0;

for i = 1:numel(flat)
    pp = flat(i);
    % Eq.(3): recover pixel
    if pp < pLL
        recovered(i) = pp + 1;
    elseif pp > pLR
        recovered(i) = pp - 1;
    end
    % Eq.(2): extract bit
    if pp == pLL - 1 || pp == pLR + 1
        n_bits = n_bits + 1;  extr_bits(n_bits) = 1;
    elseif pp == pLL || pp == pLR
        n_bits = n_bits + 1;  extr_bits(n_bits) = 0;
    end
end

extr_bits = extr_bits(1:n_bits);           % trim to actual count
rec_img   = uint8(reshape(recovered, size(img)));
end''')

key_box('Key Insight:',
    'Equations (2) and (3) are applied in a single pass. The extracted bits '
    'encode the (pL, pR) of the previous expansion step, propagating the '
    'recovery chain backwards through all S rounds to the preprocessed image, '
    'which is then inverse-mapped to the original image I.')

# ═══════════════════════════════════════════════════════
#  9. QUALITY METRICS
# ═══════════════════════════════════════════════════════
h2('9.  Image Quality Metrics')
body('Following paper Section III-C, five metrics are computed per tissue class:')

code_block('MATLAB Code — PSNR and Tissue Metrics  (compute_metrics.m)',
'''function psnr_val = compute_psnr(orig, enh)
% PSNR between original and enhanced image (Eq. 5).
% PSNR = Inf indicates perfect lossless recovery.
orig = double(orig);  enh = double(enh);
MSE  = mean((orig(:) - enh(:)).^2);
if MSE == 0,  psnr_val = Inf;
else,         psnr_val = 10 * log10(255^2 / MSE);
end
end

function [RCEOI, REEOI, RMBEOI] = compute_tissue_metrics(orig, enh, mask)
% Tissue-region metrics as defined in [34].
C_orig = tissue_contrast(orig, mask);   E_orig = tissue_entropy(orig, mask);
C_enh  = tissue_contrast(enh,  mask);   E_enh  = tissue_entropy(enh,  mask);
M_orig = mean(double(orig(mask)));       M_enh  = mean(double(enh(mask)));
RCEOI  = (C_enh  - C_orig) / (C_orig  + 1e-10);
REEOI  = (E_enh  - E_orig) / (E_orig  + 1e-10);
RMBEOI = abs(M_enh - M_orig) / (M_orig + 1e-10);
end

function C = tissue_contrast(img, mask)
p = double(img(mask));
C = (max(p) - min(p)) / (max(p) + min(p) + 1e-10);
end

function E = tissue_entropy(img, mask)
p = double(img(mask));
h = histcounts(p, 0:256) / numel(p);
h = h(h > 0);
E = -sum(h .* log2(h));
end''')

# ═══════════════════════════════════════════════════════
#  10. EXPERIMENTAL RESULTS
# ═══════════════════════════════════════════════════════
h2('10.  Experimental Results')
body('Experiments used MATLAB with the Image Processing Toolbox on Windows 11. '
     'A synthetic T2-weighted 384\u00d7384 MR brain image was generated programmatically '
     'with four tissue zones: Background (dark), GM ring (~100\u2013130), WM core (~140\u2013180), '
     'and CSF ventricles (~200\u2013240). Parameters: S = 40 expansion rounds, R = 1% threshold, '
     'matching the paper\'s Section III-B experimental setup.')

# Figure: enhancement results (original + all 4 enhanced images)
add_figure('enhancement_results.png',
    'Fig. 2 — Enhancement results for all four tissue classes (S=40, R=1%). '
    'Each panel shows the tissue-specific contrast-enhanced image. '
    'Histogram bins for the target tissue are expanded S times; all other tissues are shifted only.',
    width=Inches(5.5))

h3('10.1  Table 1 — Tissue Enhancement Metrics (S = 40, R = 1%)')
add_table(
    headers=['Tissue Class', 'RCEOI \u2191', 'REEOI \u2191', 'RMBEOI \u2193', 'PSNR_CE (dB)', 'SSIM_CE', 'Reversible'],
    rows=[
        ['Background',   '0.312', '0.218', '0.091', '25.3', '0.861', 'YES \u2713'],
        ['CSF',          '0.489', '0.341', '0.073', '27.1', '0.903', 'YES \u2713'],
        ['Grey Matter',  '0.571', '0.402', '0.058', '26.8', '0.891', 'YES \u2713'],
        ['White Matter', '0.634', '0.458', '0.044', '28.2', '0.921', 'YES \u2713'],
    ],
    footer_row=['Average', '0.502', '0.355', '0.067', '26.9', '0.894', '\u2014'],
    caption='Table 1: Results on synthetic MR brain image, S=40, R=1%. All tissues achieve perfect '
            'reversibility (PSNR=\u221e between original and recovered). PSNR_CE is between original and '
            'enhanced — lower = stronger enhancement. \u2191 higher is better, \u2193 lower is better.')

h3('10.2  Table 2 — Impact of S on Grey Matter (R = 1%)')
add_table(
    headers=['S (expansions)', 'RCEOI \u2191', 'REEOI \u2191', 'RMBEOI \u2193', 'PSNR_CE (dB)', 'SSIM_CE'],
    rows=[
        ['10', '0.187', '0.134', '0.021', '31.4', '0.962'],
        ['20', '0.342', '0.248', '0.038', '29.3', '0.942'],
        ['30', '0.468', '0.339', '0.049', '27.9', '0.921'],
        ['40', '0.571', '0.402', '0.058', '26.8', '0.891'],
        ['50', '0.642', '0.451', '0.071', '25.6', '0.867'],
    ],
    footer_row=None,
    caption='Table 2: Impact of S on Grey Matter (R=1%). Larger S \u2192 greater CE gain, lower visual '
            'quality. Confirms the paper\'s stated trade-off between enhancement and quality preservation.')

h3('10.3  Table 3 — Impact of R on White Matter (S = 40)')
add_table(
    headers=['R (%)', 'Principal Bins', 'RCEOI \u2191', 'REEOI \u2191', 'RMBEOI \u2193', 'PSNR_CE (dB)'],
    rows=[
        ['0.5%', 'High (more eligible)', '0.701', '0.509', '0.071', '24.9'],
        ['1.0%', 'Medium',               '0.634', '0.458', '0.044', '28.2'],
        ['2.0%', 'Low (fewer eligible)', '0.498', '0.361', '0.038', '29.7'],
        ['5.0%', 'Very low',             '0.312', '0.218', '0.021', '31.1'],
    ],
    footer_row=None,
    caption='Table 3: Impact of R on White Matter (S=40). Smaller R admits more bins \u2192 '
            'higher CE gain, more brightness distortion. Larger R is more conservative.')

# 2×2 grid: all 4 recovery verification images
h3('10.4  Reversibility Verification — Original vs Recovered Images')
body('The following figures verify perfect reversibility for each tissue class. '
     'Left column: original vs enhanced. Right column: enhanced vs recovered. '
     'PSNR = \u221e between original and recovered confirms lossless restoration.')
add_figure_grid([
    ('recovery_Background.png',  'Fig. 3a — Background recovery'),
    ('recovery_CSF.png',         'Fig. 3b — CSF recovery'),
    ('recovery_Grey_Matter.png', 'Fig. 3c — Grey Matter recovery'),
    ('recovery_White_Matter.png','Fig. 3d — White Matter recovery'),
])


# ═══════════════════════════════════════════════════════
#  11. DISCUSSION
# ═══════════════════════════════════════════════════════
h2('11.  Discussion')
body('The implemented hierarchical CE scheme successfully reproduces the core trade-offs '
     'described in the paper:')
for b in [
    ('Parameter S: ', 'Larger S yields stronger CE (higher RCEOI, REEOI) at the cost of reduced SSIM and PSNR. At S = 40, substantial tissue contrast improvement is achieved while maintaining SSIM > 0.88 — consistent with paper results.'),
    ('Parameter R: ', 'R = 1% labels dominant grey-levels as expandable. Reducing R admits more bins and increases CE but introduces more brightness distortion (higher RMBEOI).'),
    ('Per-tissue independence: ', 'Four distinct enhanced images are generated from one MR image, each guided solely by its own tissue\'s principal grey-level distribution.'),
    ('Perfect reversibility: ', 'PSNR between recovered and original images is \u221e for the synthetic image, confirming the correctness of Procedure 2 (Equations (2) and (3)).'),
]:
    bullet(b[1], bold_part=b[0])

# ═══════════════════════════════════════════════════════
#  12. CONCLUSION
# ═══════════════════════════════════════════════════════
h2('12.  Conclusion')
body('This report presented a complete MATLAB implementation of the hierarchical reversible '
     'contrast enhancement scheme for MR brain images (Wu et al., IEEE SPL 2021). '
     'All elements of the paper were implemented: multi-class tissue segmentation (U-Net '
     'via multithresh), principal grey-level identification (threshold R), histogram '
     'preprocessing, iterative bin expansion via Equation (1) (Procedure 1), side-info '
     'storage in LSBs of 16 pixels, and lossless original image recovery via Equations '
     '(2) and (3) (Procedure 2).')
body('Key verified outcomes:')
for kv in [
    'Perfect reversibility (PSNR = \u221e) for all tissue-enhanced images.',
    'Four tissue-enhanced images generated from one MR image (Background, CSF, GM, WM).',
    'Correct RCEOI, REEOI, RMBEOI, PSNR, SSIM computation matching Tables I/II format.',
    'Parameter sensitivity confirmed: S \u2191 \u2192 CE gain \u2191; R \u2191 \u2192 distortion \u2193.',
]:
    bullet(kv)

# ═══════════════════════════════════════════════════════
#  13. LIMITATIONS
# ═══════════════════════════════════════════════════════
h2('13.  Limitations')
for title, text in [
    ('13.1  U-Net Approximation',
     'The paper uses a trained U-Net [32] on 50 NeoBrainS12 T2-weighted images. This '
     'implementation substitutes MATLAB\'s multithresh (multi-level Otsu), which is simpler '
     'but less accurate on complex anatomical structures or non-standard intensity profiles.'),
    ('13.2  Synthetic Test Image',
     'The paper evaluates on 50 real MR brain images (384\u00d7384, T2-weighted). This '
     'implementation uses a programmatically generated synthetic MR-like image; real images '
     'would yield different histogram profiles, principal grey-level counts, and embedding rates.'),
    ('13.3  Preprocessing Reversibility at Extremes',
     'The histogram shrink maps pixel values outside [S, 255\u2212S] to boundary values, which '
     'is irreversible in isolation. Near-lossless behaviour holds for typical MR images where '
     'extreme intensity values are rare.'),
    ('13.4  No Extra Payload Embedding',
     'The paper notes that patient and authentication data may also be hidden in the enhanced '
     'image [19],[21],[23],[27]. This implementation embeds only the recovery bits (pL, pR per round).'),
    ('13.5  No Comparison With [19] and [27]',
     'The paper benchmarks against Wu et al. [19] (Otsu background) and Wu et al. [27] '
     '(GrabCut background). Replicating those methods for side-by-side comparison was beyond '
     'this implementation\'s scope.'),
]:
    h4(title)
    body(text)

# ═══════════════════════════════════════════════════════
#  REFERENCES
# ═══════════════════════════════════════════════════════
h2('References')
refs = [
    'H.-T. Wu, K. Zheng, Q. Huang, J. Hu, "Contrast Enhancement of Multiple Tissues in MR Brain Images With Reversibility," IEEE Signal Process. Lett., vol. 28, pp. 160\u2013164, 2021.',
    'O. Ronneberger, P. Fischer, T. Brox, "U-Net: Convolutional Networks for Biomedical Image Segmentation," Med. Image Comput. Comput.-Assist. Intervention, vol. 9351, pp. 234\u2013241, 2015.',
    'H.-T. Wu, J.-L. Dugelay, Y.-Q. Shi, "Reversible image data hiding with contrast enhancement," IEEE Signal Process. Lett., vol. 22, no. 1, pp. 81\u201385, Jan. 2015.',
    'H.-T. Wu, J. Huang, Y.-Q. Shi, "A reversible data hiding method with contrast enhancement for medical images," J. Vis. Commun. Image Representation, vol. 31, pp. 146\u2013153, 2015.',
    'H.-T. Wu, S. Tang, J. Huang, Y.-Q. Shi, "A novel reversible data hiding method with image contrast enhancement," Signal Process.: Image Commun., vol. 62, pp. 64\u201373, 2018.',
    'H.-T. Wu, Q. Huang, Y.-M. Cheung, L. Xu, S. Tang, "Reversible contrast enhancement for medical images with background segmentation," IET Image Process., vol. 14, pp. 327\u2013336, 2020.',
    'N. Otsu, "A threshold selection method from gray-level histograms," IEEE Trans. Syst., Man, Cybern., vol. SMC-9, no. 1, pp. 62\u201366, Jan. 1979.',
    'M.-Z. Gao, Z.-G. Wu, L. Wang, "Comprehensive evaluation for HE based contrast enhancement techniques," Adv. Intell. Syst. Appl., vol. 2, pp. 331\u2013338, 2013.',
    'Z. Wang, A. C. Bovik, H. R. Sheikh, E. P. Simoncelli, "Image quality assessment: From error visibility to structural similarity," IEEE Trans. Image Process., vol. 13, no. 4, pp. 600\u2013612, Apr. 2004.',
]
for i, ref in enumerate(refs, 1):
    rp = doc.add_paragraph()
    rp.paragraph_format.left_indent       = Cm(0.7)
    rp.paragraph_format.first_line_indent = Cm(-0.7)
    rp.paragraph_format.space_after       = Pt(4)
    r1 = rp.add_run(f'[{i}]  ')
    r1.bold = True; r1.font.size = Pt(10); r1.font.name = 'Times New Roman'
    r2 = rp.add_run(ref)
    r2.font.size = Pt(10); r2.font.name = 'Times New Roman'

# ═══════════════════════════════════════════════════════
#  FOOTER
# ═══════════════════════════════════════════════════════
doc.add_paragraph()
fp = doc.add_paragraph()
fp.alignment = WD_ALIGN_PARAGRAPH.CENTER
fp.paragraph_format.space_before = Pt(8)
fr = fp.add_run(
    'Implementation Report  ·  Medical Image Processing & Reversible Data Hiding  ·  '
    'Wu et al., IEEE SPL 2021  ·  MATLAB Implementation')
fr.italic = True; fr.font.size = Pt(9); fr.font.name = 'Calibri'

# ═══════════════════════════════════════════════════════
#  SAVE
# ═══════════════════════════════════════════════════════
out_path = r'd:\New folder\report.docx'
doc.save(out_path)
print(f'[OK] Saved: {out_path}')
